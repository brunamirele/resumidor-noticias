import os
from docx import Document
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
import requests

# === 1. Funções para converter e extrair notícias ===
def extrair_noticias_heading1(caminho_docx):
    doc = Document(caminho_docx)
    noticias = {}
    atual = []
    contador = 0
    em_noticia = False

    for p in doc.paragraphs:
        estilo = p.style.name.lower()
        texto = p.text.strip()
        if not texto:
            continue

        if estilo == "heading 1":
            if atual:
                contador += 1
                noticias[f"noticia{contador}"] = "\n".join(atual).strip()
                atual = []
            em_noticia = True
            atual.append(texto)
        elif em_noticia:
            atual.append(texto)

    if atual:
        contador += 1
        noticias[f"noticia{contador}"] = "\n".join(atual).strip()

    print(f"Total de notícias encontradas: {len(noticias)}")
    return noticias

def processar_arquivo(caminho_docx):
    if not caminho_docx.lower().endswith('.docx'):
        raise ValueError("O arquivo precisa ser .docx")
    
    print("Extraindo notícias baseadas em Heading 1...")
    noticias = extrair_noticias_heading1(caminho_docx)
    
    return noticias


# === 2. Inicializar ambiente e chain de resumo ===
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")

prompt = ChatPromptTemplate.from_template(
    "Resuma a notícia: {noticia} em até 100 palavras (não ultrapasse 100 palavras)."
)
chain = prompt | ChatOpenAI() | StrOutputParser()

# === 3. Salvar resumos em arquivo .docx ===
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import requests

# 🔍 Função de busca no Google
def buscar_link_google(titulo, veiculo):
    api_key = os.getenv("GOOGLE_SEARCH_API_KEY")
    cx = os.getenv("GOOGLE_CX")
    
    # Inclui tanto título quanto veículo para maior precisão (sem aspas)
    query = f'"{titulo}"' #{veiculo}'
    
    url = f"https://www.googleapis.com/customsearch/v1?key={api_key}&cx={cx}&q={query}&dateRestrict=d30"

    try:
        res = requests.get(url)
        data = res.json()

        total = data.get("searchInformation", {}).get("totalResults")
        print(f"[DEBUG] totalResults: {total}")

        if "items" in data and data["items"]:
            primeiro_link = data["items"][0]["link"]
            print(f"[DEBUG] Link retornado: {primeiro_link}")
            return primeiro_link
        else:
            print("[DEBUG] Nenhum item encontrado na resposta.")
    except Exception as e:
        print(f"[ERRO] Falha ao buscar no Google: {e}")
    
    return None

# 🔗 Função para inserir link clicável
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Estiliza o link (azul e sublinhado)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0000FF")
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), "single")
    rPr.append(underline)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# 📝 Função principal exportando título + link + resumo
'''def exportar_resumos_para_word(noticias_dict, resumos_dict, caminho_saida='resumos.docx'):
    doc = Document()
    doc.add_heading('Resumos das Notícias', level=1)

    for i in range(1, len(noticias_dict) + 1):
        noticia_key = f'noticia{i}'
        resumo_key = f'resumo{i}'

        noticia = noticias_dict.get(noticia_key, '')
        resumo = resumos_dict.get(resumo_key, '[Resumo não disponível]')

        linhas = noticia.split('\n')
        titulo = linhas[0] if len(linhas) > 0 else '[Título não encontrado]'
        veiculo = linhas[2] if len(linhas) > 2 else '[Veículo não identificado]'

        link = buscar_link_google(titulo, veiculo)

        # Adiciona o título
        doc.add_heading(f'{i:02d}. {titulo}', level=2)

        # Adiciona o link (caso exista)
        if link and link.startswith("http"):
            p_link = doc.add_paragraph()
            run = p_link.add_run(link)
            run.font.size = Pt(9)
            run.font.underline = True  # Mantém sublinhado (se quiser tirar, remova esta linha)

        # Adiciona o resumo
        p = doc.add_paragraph(resumo)
        p.style.font.size = Pt(11)

    doc.save(caminho_saida)
    print(f"\nArquivo Word exportado com sucesso para: {os.path.abspath(caminho_saida)}")'''

from docx.shared import Pt, RGBColor

def exportar_resumos_para_word(noticias_dict, resumos_dict, caminho_saida='resumos.docx'):
    doc = Document()
    
    # Título principal do documento
    doc.add_heading('Resumos das Notícias', level=1)

    for i in range(1, len(noticias_dict) + 1):
        noticia_key = f'noticia{i}'
        resumo_key = f'resumo{i}'

        noticia = noticias_dict.get(noticia_key, '')
        resumo = resumos_dict.get(resumo_key, '[Resumo não disponível]')

        linhas = noticia.split('\n')
        titulo = linhas[0] if len(linhas) > 0 else '[Título não encontrado]'
        veiculo = linhas[2] if len(linhas) > 2 else ''

        # Buscar link no Google
        link = buscar_link_google(titulo, veiculo)

        # 👉 Adicionar título como run para aplicar cor preta
        p_titulo = doc.add_paragraph()
        run_titulo = p_titulo.add_run(f'{i:02d}. {titulo}')
        run_titulo.bold = True
        run_titulo.font.size = Pt(14)
        run_titulo.font.color.rgb = RGBColor(0, 0, 0)  # PRETO

        # 👉 Adicionar link logo abaixo (como texto simples)
        if link and link.startswith("http"):
            p_link = doc.add_paragraph()
            run_link = p_link.add_run(link)
            run_link.font.size = Pt(9)
            run_link.font.color.rgb = RGBColor(0, 0, 255)  # azul claro padrão de link, opcional
            run_link.underline = True

        # 👉 Adicionar resumo
        p_resumo = doc.add_paragraph(resumo)
        p_resumo.style.font.size = Pt(11)

    # Salvar
    doc.save(caminho_saida)
    print(f"\nArquivo Word exportado com sucesso para: {os.path.abspath(caminho_saida)}")


# === 4. Executar todo o processo: leitura, resumo e exibição ===
def resumir_noticias(noticias_dict):
    resumos = {}
    for chave, noticia in noticias_dict.items():
        try:
            resumo = chain.invoke({'noticia': noticia})
            resumos[chave.replace("noticia", "resumo")] = resumo
        except Exception as e:
            resumos[chave.replace("noticia", "resumo")] = f"[ERRO] {e}"
    return resumos
